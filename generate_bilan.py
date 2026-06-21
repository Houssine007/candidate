# -*- coding: utf-8 -*-
"""
Génère le BILAN TECHNIQUE & FONCTIONNEL de RecruitPRO (.docx).
Document de préparation à la soutenance : cartographie complète, page par page,
de tous les services et rôles, workflows, endpoints, puis critique exhaustive.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Styles de base ───────────────────────────────────────────────────────────
NORMAL = doc.styles["Normal"]
NORMAL.font.name = "Calibri"
NORMAL.font.size = Pt(10.5)

ACCENT = RGBColor(0x0F, 0x6F, 0x5C)      # vert canard
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x2A, 0x2A)

def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(18)
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = DARK
    r.font.size = Pt(14)
    return p

def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.color.rgb = ACCENT
    r.font.size = Pt(11.5)
    return p

def para(text, italic=False, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], "0F6F5C")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    return t

def spacer():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════════════════════
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BILAN TECHNIQUE & FONCTIONNEL"); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = ACCENT
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RecruitPRO — Plateforme intégrée de gestion des compétences"); r.font.size = Pt(15); r.font.color.rgb = DARK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Recrutement (ATS) · RH interne & mobilité · Apprentissage (LMS)"); r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY
spacer(); spacer()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Document de préparation à la soutenance"); r.font.size = Pt(13); r.bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("« Connaître les failles avant la puissance »"); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cartographie de A à Z : architecture, pages, fonctionnalités, endpoints,\nworkflows, expérience utilisateur, conception, et critique complète.")
r.font.size = Pt(10.5); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Généré le " + datetime.date.today().strftime("%d/%m/%Y")); r.font.size = Pt(9); r.font.color.rgb = GREY
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE (manuel)
# ══════════════════════════════════════════════════════════════════════════════
h1("Sommaire")
toc = [
 "1. Résumé exécutif",
 "2. Vision, problématique et positionnement",
 "3. Architecture générale du système",
 "4. Authentification unifiée (SSO) et sécurité",
 "5. Modèle de données",
 "6. Le moteur de matching (cœur algorithmique)",
 "7. Cartographie page par page — Frontend RH",
 "8. Cartographie page par page — LMS (Academy)",
 "9. Inventaire des endpoints (API)",
 "10. Workflows clés de bout en bout",
 "11. Conception : choix d'ingénierie et justifications",
 "12. CRITIQUE COMPLÈTE — les failles",
 "13. Points forts et éléments différenciants",
 "14. Zones de flou / à clarifier avant la soutenance",
 "15. Questions probables du jury et éléments de réponse",
 "16. Feuille de route d'amélioration",
]
for t in toc: para(t)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. RÉSUMÉ EXÉCUTIF
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Résumé exécutif")
para("RecruitPRO est une plateforme RH « tout-en-un » qui réunit trois domaines habituellement "
     "traités par des outils séparés : le recrutement externe (ATS), la gestion RH interne "
     "(employés, organigramme, mobilité interne) et la formation (LMS). Le fil conducteur, et "
     "l'ambition intellectuelle du projet, est le référentiel de compétences vivant : une même "
     "compétence circule entre le candidat, l'employé, l'offre, la formation et le pilotage "
     "prévisionnel des effectifs (GPEC), et se met à jour automatiquement quand le collaborateur "
     "progresse.")
para("Le système est composé de trois applications indépendantes qui partagent une seule identité "
     "JWT (authentification unique) :")
table(
 ["Application", "Rôle", "Stack", "Port", "Stockage"],
 [
  ["Backend API", "Cerveau métier (matching, RBAC, GPEC)", "FastAPI + SQLAlchemy + Alembic (Python)", "8000", "PostgreSQL"],
  ["Frontend RH", "Interfaces recruteur / candidat / employé", "Next.js 15, TypeScript, Tailwind v3, Zustand", "3000", "—"],
  ["LMS (Academy)", "Cours, examens, validation de compétences", "Next.js 16, TypeScript, Tailwind v4, Mongoose", "3001", "MongoDB"],
 ],
 widths=[1.3, 2.4, 2.3, 0.5, 1.0]
)
spacer()
para("Différenciateur central — la « boucle vivante » :", bold=True)
para("Un recruteur (ou un manager) détecte un écart de compétence → assigne une formation LMS à "
     "l'employé → l'employé suit le cours et réussit l'examen final → le LMS notifie le backend "
     "RH → la compétence est relevée (jamais abaissée) dans le profil candidat ET le profil "
     "employé → la cartographie GPEC se met à jour. Cette boucle est réellement implémentée "
     "(double-écriture côté backend), ce qui distingue le projet d'une simple juxtaposition "
     "d'outils.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. VISION
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Vision, problématique et positionnement")
h2("2.1 Problématique")
para("Les organisations savent recruter et savent former, mais peinent à relier ces deux gestes : "
     "le recrutement raisonne sur les compétences détenues aujourd'hui (le « savoir-faire »), "
     "tandis que la gestion des talents devrait raisonner sur la capacité à évoluer (le "
     "« savoir-évoluer »). RecruitPRO tente de faire tenir ces deux logiques dans un même "
     "référentiel, pour piloter la compétence dans la durée plutôt que de la photographier à "
     "l'instant de l'embauche.")
h2("2.2 Les trois mécanismes")
bullet("le référentiel ROME comme langage commun : chaque compétence est rattachée à un code "
       "métier (France Travail / ROME), ce qui permet de mesurer la proximité entre métiers "
       "(adjacence) et donc le potentiel d'évolution.", "Sémantique partagée — ")
bullet("deux axes de lecture d'un profil : le « fit » (adéquation au besoin actuel) et le "
       "« potentiel » (capacité à combler les écarts par la formation).", "Double évaluation — ")
bullet("la compétence validée par la formation remonte automatiquement dans les profils et la "
       "GPEC : le référentiel n'est pas figé, il « vit ».", "Boucle de rétroaction — ")
h2("2.3 Positionnement (contexte mémoire)")
para("Le projet est présenté comme une plateforme stratégique de gestion des compétences "
     "(orientation GPEC / capital humain), au-delà d'un simple ATS. L'argument de soutenance "
     "n'est donc pas « j'ai codé un site de recrutement » mais « j'ai conçu un dispositif qui "
     "fait dialoguer recrutement, mobilité et formation autour d'un référentiel de compétences "
     "unique et évolutif ».")

# ══════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Architecture générale du système")
h2("3.1 Vue d'ensemble")
para("Trois applications, trois bases de données logiques, un secret JWT commun. Le frontend RH "
     "est le seul point d'entrée (login) ; le LMS n'a pas de page de connexion propre.")
para("Schéma logique (textuel) :", bold=True)
para("Navigateur ⇄ Frontend RH (Next.js :3000) ⇄ Backend API (FastAPI :8000) ⇄ PostgreSQL", italic=True)
para("Frontend RH ── ouverture SSO (?token=) ──▶ LMS (Next.js :3001) ⇄ MongoDB", italic=True)
para("Backend API ── token de service (rôle ADMIN, 5 min) ──▶ LMS API (proxy cours/inscriptions)", italic=True)
para("LMS ── POST /api/lms/course-completed (à la réussite) ──▶ Backend API (relève la compétence)", italic=True)
h2("3.2 Backend — architecture en couches")
para("Application FastAPI monolithique et organisée en couches :")
bullet("app/api/ : 17 routers montés sous /api/* (auth, candidates, jobs, skills, catalog, "
       "applications, companies, recruiters, users, organization, employees, roles, lms, "
       "internal_mobility, trainings, gpec).", "Routes — ")
bullet("app/models/ : modèles SQLAlchemy (User, Candidate, Employee, Job, Skill, Application, "
       "InternalPosition, Training, OrgUnit, JobStandard, WorkforceTarget, Permission/InternalRole…).", "Modèles — ")
bullet("app/services/ : logique métier (matching, forecast, profile_analysis, completeness, "
       "cv_service, rome_api / france_travail_api, scoring).", "Services — ")
bullet("app/core/ : config, database, security (bcrypt), permissions (dépendance RBAC).", "Infra — ")
h2("3.3 Multi-tenant")
para("Modèle « base partagée / schéma partagé ». Les tables rattachées à une entreprise portent "
     "un company_id (jobs, employees, org_units, internal_roles, internal_positions, trainings, "
     "workforce_targets). Les APIs filtrent par l'entreprise du recruteur connecté. Choix "
     "assumé : les candidats et les compétences sont GLOBAUX (vivier de talents partagé + "
     "référentiel de compétences commun), donc sans company_id.")
h2("3.4 RBAC à deux niveaux")
bullet("Rôle système (users.role) : ADMIN | RECRUITER | CANDIDATE | EMPLOYEE — vérifié "
       "directement dans les routes.", "Niveau 1 — ")
bullet("Rôle interne + permissions par entreprise : dépendance has_permission(\"perm\") "
       "(app/core/permissions.py). Règles de contournement : l'ADMIN global a tout ; un "
       "RECRUITER sans fiche employé (le propriétaire du tenant) a tout ; sinon la permission "
       "doit figurer dans le rôle interne de l'employé.", "Niveau 2 — ")

# ══════════════════════════════════════════════════════════════════════════════
# 4. SSO & SÉCURITÉ
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Authentification unifiée (SSO) et sécurité")
h2("4.1 Principe")
para("Les trois applications font confiance au même secret JWT HS256 (SECRET_KEY côté backend / "
     "JWT_SECRET côté LMS). Changer ce secret dans une seule application casse le SSO partout.")
h2("4.2 Déroulé du SSO")
numbered("L'utilisateur se connecte sur le frontend RH (POST /api/auth/token). Le backend renvoie "
         "un access_token JWT dont les claims sont : sub (= id utilisateur), email, full_name, "
         "role, is_instructor, company_id.")
numbered("Pour ouvrir le LMS, le RH construit une URL lmsLaunchUrl(token, path) qui ajoute "
         "?token=<jwt>. Le LMS capte le token (captureSsoToken), le persiste dans localStorage, "
         "puis nettoie l'URL (history.replaceState).")
numbered("Les routes API du LMS vérifient le token avec la librairie jose contre le secret "
         "partagé (lms/lib/auth.ts).")
numbered("Pour les appels backend → LMS, le backend forge un token de service éphémère "
         "(_get_service_token : sub=\"0\", role ADMIN, 5 min) et proxifie vers LMS_API_URL.")
numbered("À la réussite d'un cours, le LMS appelle POST /api/lms/course-completed qui relève la "
         "compétence correspondante (sans jamais l'abaisser).")
h2("4.3 Autorisation côté LMS (instructeurs)")
para("Toutes les routes /api/instructor/* passent par authorizeInstructor(req) : le JWT RH est "
     "vérifié, et l'accès est accordé si is_instructor=true OU rôle système ADMIN. La propriété "
     "d'un cours/catégorie est scoping par l'id utilisateur RH numérique (Course.instructorId et "
     "Category.instructorId sont des Number = id Postgres, pas des ObjectId Mongo). "
     "/api/auth/me mappe is_instructor → role:'instructor' uniquement pour les gardes côté "
     "client.")
para("Hachage des mots de passe : bcrypt (passlib). Aucun mot de passe en clair stocké.", )

# ══════════════════════════════════════════════════════════════════════════════
# 5. MODÈLE DE DONNÉES
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Modèle de données")
h2("5.1 Entités principales (PostgreSQL)")
table(
 ["Entité", "Rôle", "Champs / relations notables"],
 [
  ["User", "Identité & rôle système", "email, password (bcrypt), role (enum), is_instructor ; 1-1 vers Recruiter/Candidate/Employee"],
  ["Candidate", "Profil du vivier global", "bio, cv_text, cv_url, formations, certifications, projects (JSON), links (JSON), onboarding_step, profile_completeness_score ; skills"],
  ["CandidateSkill", "Compétence d'un candidat", "level (1-4), years_experience ; FK skill"],
  ["Employee", "Collaborateur d'une entreprise", "company_id, org_unit_id, internal_role_id, manager_id (auto-réf.), job_title ; skills"],
  ["EmployeeSkill", "Compétence d'un employé", "level (1-4), years_experience, certified ; FK skill"],
  ["Skill", "Référentiel de compétences (GLOBAL)", "name, category, rome_code"],
  ["Job", "Offre d'emploi", "company_id, recruiter_id, min_years_experience, min_education_level ; requirements"],
  ["JobRequirement", "Compétence exigée par une offre", "required_level, is_mandatory ; FK skill"],
  ["Application", "Candidature", "status (enum), cover_letter ; FK candidate / job"],
  ["InternalPosition / …Requirement / InternalApplication", "Mobilité interne", "statut OPEN/CLOSED/FILLED ; candidatures PENDING→ACCEPTED/REJECTED"],
  ["Training / TrainingSkill / TrainingEnrollment", "Catalogue formations interne", "catégorie, skills_taught (level_gained), enrollment status"],
  ["OrgUnit", "Unité de l'organigramme", "parent_id (arbre), manager_id, unit_type"],
  ["JobStandard / …Requirement", "Emploi-type de référence (ROME)", "rome_code ; min_level, is_mandatory, is_emergent"],
  ["WorkforceTarget", "Cible d'effectif (GPEC prévisionnelle)", "job_standard_id, org_unit_id, target_headcount, horizon, priority"],
  ["Permission / InternalRole / role_permissions", "RBAC interne", "M-N rôles↔permissions par entreprise"],
  ["Evaluation", "Évaluation annuelle (modèle présent)", "rating, forces, axes ; peu/pas exposé en UI"],
 ],
 widths=[1.8, 2.0, 3.0]
)
h2("5.2 Données LMS (MongoDB)")
para("Modèles Mongoose : Course, Module, Section, Quiz, Question, Answer, Enrollment, Progress, "
     "Category. Points clés : Course.skillId (= id de Skill RH) et Course.skillLevel (niveau "
     "atteint à la réussite) font le pont entre le LMS et le référentiel RH ; Enrollment "
     "(employeeId, status assigned/in_progress/completed/abandoned, finalScore, progress).")
para("Niveaux de compétence : entiers 1 à 4 partout (candidat, employé, exigence, formation) — "
     "convention homogène.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. MATCHING
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Le moteur de matching (cœur algorithmique)")
para("Fichier : backend/app/services/matching.py. Deux axes indépendants sont calculés pour "
     "chaque couple (profil, besoin).")
h2("6.1 Fit score — adéquation au besoin actuel")
bullet("Compétences : 60 %. Chaque exigence est pondérée (obligatoire ×2, sinon ×1). Le ratio "
       "niveau_candidat / niveau_requis est plafonné à 1,2. Pénalité de −30 % par compétence "
       "obligatoire totalement absente (facteur 0,7^n).", "")
bullet("Expérience : 25 % (années candidat / années requises, plafonné à 100 %).", "")
bullet("Éducation : 15 % (niveau candidat / niveau requis, plafonné à 100 %).", "")
para("Score final = 0,60×compétences + 0,25×expérience + 0,15×éducation.", italic=True)
h2("6.2 Potential score — capacité à évoluer par la formation")
para("Calculé uniquement s'il existe des écarts (sinon None). 100 % de signaux factuels positifs, "
     "aucun malus. Pour chaque écart de compétence, on mesure l'adjacence ROME avec ce que le "
     "candidat maîtrise déjà :")
bullet("même sous-domaine ROME (ex. M1805 = M1805) → poids 1,0 (pont « exact »)")
bullet("même grand domaine ROME (ex. M18xx) → poids 0,6 (pont « domaine »)")
bullet("sans rapport → poids 0,3 (à former)")
para("potential = (moyenne des poids) × 80 + bonus certifications (+12 si le champ certifications "
     "est rempli). Plafonné à 100.", italic=True)
h2("6.3 Recommandation")
para("Étiquette dérivée des deux axes : STRONG_FIT si fit ≥ 70 ; sinon POTENTIAL si potentiel ≥ "
     "60 ; sinon WEAK_FIT. Affichée comme « Prêt » / « Haut potentiel » / « Écart important ».")
h2("6.4 Réutilisation pour la GPEC")
para("Le service forecast.py réutilise intégralement ce moteur : chaque EMPLOYÉ est traité comme "
     "un « candidat » et chaque emploi-type cible (JobStandard) comme une « offre ». On en déduit "
     "les qualifiés (tous les obligatoires au niveau), les potentiels (formables, potentiel ≥ 60), "
     "l'écart (cible − qualifiés) et un statut critique/partiel/couvert. Aucune logique de "
     "scoring n'est dupliquée.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. FRONTEND RH — PAGE PAR PAGE
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Cartographie page par page — Frontend RH")
para("Organisation App Router : dashboard/{recruiter,candidate,employee}. État d'auth dans le "
     "store Zustand persistant (recruitpro-auth). Tous les appels et types sont centralisés dans "
     "src/lib/api.ts. Thème clair/sombre via next-themes. La sidebar recruteur est mutualisée "
     "(components/recruiter-sidebar.tsx).")

def page_block(titre, route, role, but, contenu, endpoints, ux=None):
    h3(titre)
    para("Route : " + route + "   |   Rôle : " + role, italic=True, color=GREY, size=9)
    para("Objectif : " + but, bold=True)
    para(contenu)
    if ux:
        para("UX / navigation : " + ux)
    para("Endpoints / appels : " + endpoints, italic=True, color=GREY, size=9)

h2("7.1 Pages publiques")
page_block(
 "Landing (accueil marketing)", "/", "public",
 "Vitrine du produit et amorce de conversion.",
 "Navbar verre dépoli, hero « Transformez votre Capital Humain », section « Dernières "
 "Opportunités » qui charge les 6 dernières offres, grille « bento » de features, CTA, footer. "
 "Un candidat connecté peut postuler directement depuis une carte d'offre (bouton « Postuler »).",
 "getLatestJobs(6) ; applyToJob(jobId) si CANDIDATE connecté.",
 "Si non connecté → redirige vers /login ; si rôle ≠ CANDIDATE → alerte native. Plusieurs "
 "chiffres sont décoratifs/codés en dur (« 99.8 % », « +500 entreprises », « Posté il y a 2j »).")
page_block(
 "Connexion", "/login", "public",
 "Point d'entrée unique de toute la plateforme (y compris LMS).",
 "Écran deux colonnes (branding + formulaire). Champs e-mail / mot de passe pré-remplis avec un "
 "compte de démo. Après login, redirection selon le rôle : RECRUITER/ADMIN → dashboard "
 "recruteur, EMPLOYEE → espace employé, CANDIDATE → espace candidat.",
 "login(email, password) → POST /api/auth/token puis GET /api/auth/me.")
page_block(
 "Inscription", "/signup", "public",
 "Créer un compte candidat ou recruteur.",
 "Formulaire nom complet / e-mail / mot de passe + choix du rôle. Un nouveau candidat est "
 "redirigé vers l'onboarding ; un recruteur vers son dashboard.",
 "register(fullName, email, password, role) → POST /api/auth/register.")

h2("7.2 Espace candidat")
page_block(
 "Tableau de bord candidat", "/dashboard/candidate", "CANDIDATE",
 "Suivre ses candidatures et optimiser son profil.",
 "Sidebar « TalentSpace » + navigation mobile. KPIs (candidatures actives, score de match moyen, "
 "visibilité). Liste des candidatures avec badge de statut (À l'étude / En revue / Finaliste / "
 "Recruté / Refusé) et score de compatibilité. Modale d'édition de profil riche : upload CV "
 "(analyse IA), édition des compétences (niveaux 1-4), recherche dans le catalogue ROME. Bloc "
 "« Analyse de votre profil » qui regroupe les compétences par domaine métier (fort / potentiel "
 "/ à développer) et suggère des compétences à acquérir.",
 "getMyApplications, getMyProfile, getSkills, updateMyProfile, uploadCV, searchSkills, "
 "getProfileAnalysis.",
 "La carte « Visibilité : Expert » est statique. La section analyse appelle "
 "/candidates/me/profile-analysis (regroupement ROME, jamais exposé en clair).")
page_block(
 "Onboarding candidat (assistant)", "/dashboard/candidate/onboarding", "CANDIDATE (et EMPLOYEE en édition)",
 "Construire un profil complet, étape par étape, et maximiser un score de complétude.",
 "Assistant en 9 étapes : (0) upload CV avec extraction IA, (1) identité, (2) liens "
 "GitHub/portfolio/LinkedIn, (3) formation, (4) expériences, (5) projets, (6) compétences "
 "(recherche ROME, niveaux, source affichée : Déclaré / Extrait CV / Projet), (7) certifications, "
 "(8) récap avec score visuel et checklist pondérée. Sauvegarde progressive (onboarding_step). "
 "Un score < 40 % est signalé comme insuffisant pour postuler.",
 "getMyProfile, uploadCV, searchSkills, updateCandidateOnboarding (PATCH par étape).",
 "Le même écran sert à un employé qui édite son profil (retour vers /dashboard/employee).")

h2("7.3 Espace employé")
page_block(
 "Tableau de bord employé", "/dashboard/employee", "EMPLOYEE",
 "Vue de synthèse du collaborateur : profil, rôle interne, compétences, accès.",
 "Sidebar (Mon Profil, Mes Formations, Mobilité Interne, et « Outils RH » si l'employé a des "
 "permissions de gestion). KPIs (rôle interne, nombre de compétences, nombre de permissions). "
 "Bloc « Profil professionnel » réutilisant le profil candidat « emporté » lors de l'embauche "
 "(bio, parcours, diplômes, certifications, CV). Liste des compétences avec niveaux.",
 "getEmployeeMe, getUserPermissions, getMyProfile.",
 "Bloc « Mes Formations » du dashboard codé en dur (« Aucune formation assignée ») — la vraie "
 "liste est sur la page dédiée /dashboard/employee/formations.")
page_block(
 "Mes formations (employé)", "/dashboard/employee/formations", "EMPLOYEE",
 "Voir les formations assignées (catalogue interne + cours LMS) et accéder au player.",
 "Unifie deux sources : les inscriptions au catalogue interne (TrainingEnrollment) et les "
 "inscriptions LMS (Enrollment). Permet de lancer un cours LMS en SSO.",
 "getMyTrainingEnrollments, getMyLmsEnrollments.")
page_block(
 "Mobilité interne (employé)", "/dashboard/employee/mobilite", "EMPLOYEE",
 "Découvrir les postes internes ouverts et postuler.",
 "Liste des postes internes (titre, lieu, fourchette de salaire, compétences requises). Bouton "
 "« Postuler » ouvrant une modale de lettre de motivation. Statut de candidature affiché "
 "(Candidature envoyée / En revue / Entretien).",
 "getInternalPositions, getMyInternalApplications, applyToPosition.")

h2("7.4 Espace recruteur / RH")
page_block(
 "Tableau de bord recruteur", "/dashboard/recruiter", "RECRUITER / ADMIN",
 "Piloter les recrutements : offres ouvertes et talents suggérés.",
 "Sidebar mutualisée (Vue d'ensemble, Employés, Organisation, Pilotage GPEC, Formations, + "
 "bouton SSO « Académie (LMS) »). KPIs (postes ouverts, candidats matchés, interviews). Pour "
 "chaque offre : description, critères, et liste de talents suggérés (sourcing) avec score de "
 "match, indicateur de gaps, et actions (voir profil, inviter dans le pipeline, accéder au "
 "pipeline).",
 "getRecruiterJobsWithMatches (liste des offres puis /jobs/{id}/matches pour chacune), "
 "inviteCandidate.",
 "Le nom d'entreprise affiché (« chez TechCorp ») est codé en dur. L'enrichissement des matches "
 "fait un appel par offre (N+1).")
page_block(
 "Création d'offre (assistant + IA)", "/dashboard/recruiter/jobs/create", "RECRUITER / ADMIN",
 "Créer une offre structurée avec assistance ROME et génération IA.",
 "Assistant 4 étapes (Poste, Compétences, Critères, Publication). À la saisie du titre, deux "
 "aides apparaissent : (1) suggestions de fiches ROME (compétences + description pré-rédigée), "
 "(2) génération par IA (Groq) d'une description et d'une liste de compétences. Les compétences "
 "importées sont tracées par source (rome / ai / manual), avec niveau 1-4 et caractère "
 "obligatoire. Aperçu avant publication, panneau latéral de résumé et de complétude.",
 "/api/catalog/jobs/suggest (ROME), /api/catalog/jobs/suggest-ai (Groq), searchSkills, createJob.",
 "Le champ company est figé sur « TechCorp » à la création.")
page_block(
 "Création d'offre (formulaire simple)", "/dashboard/recruiter/jobs/new", "RECRUITER / ADMIN",
 "Variante simplifiée de création d'offre (formulaire classique).",
 "Formulaire titre / lieu / salaire / description + ajout manuel de compétences requises. "
 "Coexiste avec l'assistant ci-dessus (doublon fonctionnel).",
 "getSkills, createJob.")
page_block(
 "Pipeline de candidatures (Kanban)", "/dashboard/recruiter/applications/[jobId]", "RECRUITER / ADMIN",
 "Gérer le suivi des candidatures d'une offre et explorer la mobilité interne.",
 "Tableau Kanban à 5 colonnes (À examiner, En revue, Finalistes, Recrutés, Refusés). Bascule de "
 "tri Fit ↔ Potentiel. Carte candidat avec score, badge de recommandation et gaps. Modale "
 "« dossier candidat » très détaillée : cercle de score, décomposition skills/XP/éducation, "
 "potentiel, bio/parcours/formations/certifications, texte du CV, comparaison compétence par "
 "compétence, analyse des gaps avec badges de « pontabilité » (Pontable / Proche / À former). "
 "Changement de statut ; modales de confirmation pour l'embauche (titre de poste définitif) et "
 "le refus. Panneau « Mobilité interne » : employés qui pourraient combler le poste avec "
 "formation, avec assignation de cours LMS.",
 "getJob, getJobApplications, updateApplicationStatus, confirmHire (POST /hire), "
 "getInternalMatches (/jobs/{id}/internal-matches), getLMSCourses, assignCourse.",
 "Le rapprochement cours↔compétence dans la mobilité se fait par correspondance textuelle de "
 "titre (heuristique).")
page_block(
 "Profil candidat détaillé", "/dashboard/recruiter/candidates/[id]", "RECRUITER / ADMIN",
 "Consulter en profondeur un candidat et préparer un plan de formation.",
 "Fiche complète (identité, score de match avec l'offre en contexte, compétences, gaps & "
 "formations LMS suggérées pour les combler). Permet de relier directement un écart à une "
 "formation.",
 "getCandidate(id, token, jobId), getLMSCourses.")
page_block(
 "Gestion des employés", "/dashboard/recruiter/employees", "RECRUITER",
 "CRUD des collaborateurs.",
 "Liste/recherche des employés (nom, poste, unité, rôle), création (modale avec identité, poste, "
 "compétences), suppression (avec confirmation). Affiche l'unité d'appartenance.",
 "getEmployees, getInternalRoles, getOrgTree, createEmployee, deleteEmployee.")
page_block(
 "Organisation (organigramme)", "/dashboard/recruiter/organization", "RECRUITER / ADMIN",
 "Structurer l'entreprise et affecter les collaborateurs et leurs rôles de sécurité.",
 "Organigramme arborescent (unités/sous-unités) avec membres affichés en pastilles. Panneau "
 "« À assigner » (employés sans unité) et résumé des rôles de sécurité avec leur nombre de "
 "permissions. Modales : créer une unité (parent/enfant) ; assigner un collaborateur à une "
 "unité + un rôle interne.",
 "getOrgTree, getEmployees, getInternalRoles, createOrgUnit, updateEmployee.")
page_block(
 "Formations & compétences (RH)", "/dashboard/recruiter/formations", "RECRUITER",
 "Assigner des cours LMS aux employés et suivre l'avancement.",
 "Sélection d'un employé, vue de ses inscriptions/avancement, catalogue de cours LMS recherchable "
 "et assignation.",
 "getEmployees, getLMSCourses, getEmployeeEnrollments, assignCourse.")
page_block(
 "Pilotage GPEC", "/dashboard/recruiter/gpec", "RECRUITER / ADMIN",
 "Cartographier les compétences détenues, analyser les écarts, et planifier (prévisionnel).",
 "Deux vues commutables. (A) « Postes ouverts » : KPIs (effectif, compétences détenues, postes "
 "ouverts, compétences en tension), indice de couverture global avec objectif, état des lieux par "
 "domaine (niveau moyen), table d'analyse des écarts par compétence (détiennent / au niveau / "
 "statut critique-partiel-couvert / impact postes & équipes) avec bouton « Agir » qui ouvre une "
 "modale de plan d'action (choisir une formation LMS et les employés sous le niveau requis, puis "
 "assigner en masse). (B) « Prévisionnel » : couverture des emplois-types cibles par l'effectif à "
 "horizon (qualifiés / potentiels / écart), avec synthèse générée par IA.",
 "getGpecOverview, getGpecSkillGap, getCoursesForSkill, assignCourse ; getGpecForecast, "
 "getGpecForecastSynthesis, getGpecTargetPool, getGpecMobility.",
 "Cette page matérialise la boucle vivante : l'écart détecté débouche directement sur une "
 "assignation de formation.")

# ══════════════════════════════════════════════════════════════════════════════
# 8. LMS — PAGE PAR PAGE
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Cartographie page par page — LMS (Academy)")
para("Application Next.js où les routes API sont le backend. Pas de login propre : le token RH "
     "est reçu en SSO. Deux univers : l'espace étudiant (français, rebrandé « RecruitPRO "
     "Academy ») et l'espace formateur (hérité d'un template « Dar Al-Ilm », en anglais).")
page_block(
 "Catalogue de cours", "/courses", "étudiant (employé)",
 "Parcourir les cours disponibles.",
 "Liste des cours publiés.",
 "GET /api/courses.")
page_block(
 "Lecteur de cours + examen", "/courses/[id]", "étudiant (employé)",
 "Suivre le contenu et passer l'examen final qui valide une compétence.",
 "Trois modes : Contenu (modules → sections : vidéos YouTube intégrées, documents téléchargeables), "
 "Examen (QCM à réponse unique ou multiple, seuil de réussite), Résultat (score, réussite, et le "
 "cas échéant « Compétence validée au niveau N — votre profil RH a été mis à jour »). Soumission "
 "liée à l'inscription de l'utilisateur.",
 "GET /api/courses/{id}, GET /api/enrollments/me, POST /api/enrollments/{id}/submit-exam.",
 "C'est ici que se déclenche la remontée de compétence vers le RH (course-completed).")
page_block(
 "Espace formateur — Dashboard", "/instructor", "is_instructor / ADMIN",
 "Vue d'ensemble du formateur.",
 "Statistiques (total cours, publiés/brouillons, étudiants, inscriptions), actions rapides "
 "(créer un cours), liste des derniers cours.",
 "GET /api/instructor/statistics, GET /api/instructor/courses.",
 "Interface en anglais, branding « Dar Al-Ilm ». La sidebar contient des entrées Students et "
 "Analytics qui ne correspondent à aucune page existante (liens morts).")
page_block(
 "Mes cours (formateur)", "/instructor/courses", "is_instructor / ADMIN",
 "Lister et gérer ses cours.",
 "Tableau des cours avec statut publié/brouillon.",
 "GET /api/instructor/courses.")
page_block(
 "Créer un cours", "/instructor/courses/create", "is_instructor / ADMIN",
 "Créer un cours et le relier à une compétence RH.",
 "Formulaire (titre, description, catégorie, vignette, statut) + champs clés skillId (compétence "
 "RH validée) et skillLevel (niveau atteint à la réussite). C'est ce lien qui permet à l'examen "
 "réussi de relever la bonne compétence côté RH.",
 "POST /api/instructor/courses.")
page_block(
 "Éditer un cours / examen", "/instructor/courses/[id]/edit", "is_instructor / ADMIN",
 "Construire le contenu pédagogique et l'examen final.",
 "Gestion des modules, sections, et examen final (titre, description, seuil de réussite, limite "
 "de temps, questions/réponses). Publication du cours.",
 "Routes /api/instructor/courses/{id}/modules, /modules/{id}/sections, /…/final-exam, "
 "/quizzes/{id}/questions, /questions/{id}/answers, /upload.")
page_block(
 "Catégories", "/instructor/categories", "is_instructor / ADMIN",
 "Organiser les cours par catégorie.",
 "CRUD de catégories (scopées par instructeur).",
 "GET/POST /api/instructor/categories, PUT/DELETE /api/instructor/categories/{id}.")
page_block(
 "Profil / Paramètres formateur", "/instructor/profile, /instructor/settings", "is_instructor / ADMIN",
 "Réglages du compte formateur.",
 "Pages de profil et de paramètres (héritées du template).",
 "—")
page_block(
 "Dashboard LMS (hérité)", "/dashboard", "tout utilisateur LMS",
 "Page d'accueil générique du template d'origine.",
 "Affiche prénom/nom/rôle. Vestige du template : lit des champs (firstName, lastName, role "
 "student/instructor/admin) qui ne correspondent pas exactement au mapping réel du JWT RH.",
 "GET /api/auth/me.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Inventaire des endpoints (API)")
h2("9.1 Backend FastAPI (sélection représentative)")
table(
 ["Domaine", "Méthode & route", "Fonction"],
 [
  ["Auth", "POST /api/auth/register", "Inscription"],
  ["Auth", "POST /api/auth/token", "Connexion (OAuth2 password) → JWT"],
  ["Auth", "GET /api/auth/me", "Utilisateur courant"],
  ["Candidats", "GET/PUT /api/candidates/me/", "Profil du candidat connecté"],
  ["Candidats", "PATCH /api/candidates/me/onboarding/", "Sauvegarde d'étape d'onboarding"],
  ["Candidats", "POST /api/candidates/me/cv/", "Upload + analyse CV"],
  ["Candidats", "GET /api/candidates/me/profile-analysis", "Analyse de profil (domaines ROME)"],
  ["Candidats", "POST /api/candidates/discover", "Découverte de talents (offre simulée)"],
  ["Offres", "POST/GET /api/jobs/", "Créer / lister les offres"],
  ["Offres", "GET /api/jobs/{id}/matches", "Candidats correspondants (matching)"],
  ["Offres", "GET /api/jobs/{id}/internal-matches", "Employés mobilisables (mobilité)"],
  ["Candidatures", "POST /api/applications/ , /invite", "Postuler / inviter un candidat"],
  ["Candidatures", "PATCH /api/applications/{id}/status/", "Changer le statut (Kanban)"],
  ["Candidatures", "POST /api/applications/{id}/hire", "Embaucher → créer l'employé"],
  ["Catalogue", "GET /api/catalog/jobs/suggest", "Suggestions ROME (emplois-types)"],
  ["Catalogue", "GET /api/catalog/jobs/suggest-ai", "Génération IA (Groq) d'offre"],
  ["Employés", "GET/POST/PUT/DELETE /api/employees", "CRUD employés"],
  ["Employés", "GET /api/employees/orgchart/tree", "Organigramme"],
  ["Organisation", "GET /api/organization/tree", "Arbre des unités"],
  ["Rôles", "GET/POST /api/roles , /permissions", "Rôles internes & permissions"],
  ["Mobilité", "GET /internal-mobility/positions", "Postes internes ouverts"],
  ["Mobilité", "POST /internal-mobility/applications", "Candidature interne"],
  ["Formations", "GET /trainings/catalog , /my-enrollments", "Catalogue & inscriptions internes"],
  ["GPEC", "GET /api/gpec/overview", "Cartographie & écarts (postes ouverts)"],
  ["GPEC", "GET /api/gpec/skill/{id}/gap", "Employés sous le niveau requis"],
  ["GPEC", "GET/POST/DELETE /api/gpec/targets", "Cibles d'effectif (prévisionnel)"],
  ["GPEC", "GET /api/gpec/forecast , /forecast/synthesis", "Couverture prévisionnelle + synthèse IA"],
  ["GPEC", "GET /api/gpec/targets/{id}/pool , /mobility/{emp}", "Vivier interne & explication mobilité"],
  ["LMS (pont)", "POST /api/lms/course-completed", "Relève la compétence à la réussite"],
  ["LMS (pont)", "GET /api/lms/courses , /enrollments ; POST /enroll", "Proxy vers le LMS"],
 ],
 widths=[1.2, 2.9, 2.7]
)
h2("9.2 LMS (routes Next.js)")
table(
 ["Méthode & route", "Fonction"],
 [
  ["GET /api/auth/me", "Identité (mappe is_instructor → role 'instructor')"],
  ["GET /api/courses , /api/courses/{id}", "Catalogue & détail cours"],
  ["GET/POST /api/enrollments , GET /api/enrollments/me", "Inscriptions (assignation, mes cours)"],
  ["PATCH /api/enrollments/{id}/progress", "Avancement"],
  ["POST /api/enrollments/{id}/submit-exam", "Soumission examen + validation compétence"],
  ["/api/instructor/courses[/...]", "CRUD cours, modules, sections, examen final"],
  ["/api/instructor/categories[/...]", "CRUD catégories"],
  ["/api/instructor/quizzes/{id}/questions , /questions/{id}/answers", "Questions & réponses"],
  ["GET /api/instructor/statistics", "Statistiques formateur"],
  ["POST /api/instructor/upload", "Upload de fichiers pédagogiques"],
 ],
 widths=[3.6, 3.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. WORKFLOWS
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Workflows clés de bout en bout")
h2("10.1 Du candidat à l'employé (recrutement)")
numbered("Le candidat s'inscrit, complète l'onboarding (ou importe son CV), déclare ses "
         "compétences (niveaux 1-4).")
numbered("Il postule à une offre (ou est invité par le recruteur). Le matching calcule fit + "
         "potentiel + gaps.")
numbered("Le recruteur suit la candidature dans le Kanban, consulte le dossier, change le statut.")
numbered("À l'acceptation (« Confirmer l'intégration », titre de poste définitif), le backend "
         "passe la candidature à ACCEPTED puis promeut le candidat en employé : bascule du rôle "
         "système CANDIDATE → EMPLOYEE, création de la fiche Employee dans l'entreprise, copie des "
         "compétences. Le profil candidat (bio, CV, compétences) est « emporté » car lié au même "
         "user_id. Opération idempotente.")
h2("10.2 La boucle vivante (formation → compétence → GPEC)")
numbered("Le RH (dashboard GPEC, pipeline mobilité, ou page formations) détecte un écart et "
         "assigne un cours LMS à l'employé (POST /api/lms/enroll → LMS).")
numbered("L'employé suit le cours dans le LMS (SSO) et passe l'examen final.")
numbered("Si réussi, le LMS appelle POST /api/lms/course-completed avec skill_id et skill_level.")
numbered("Le backend relève le niveau de la compétence — jamais à la baisse — dans le profil "
         "candidat ET dans le(s) profil(s) employé (double-écriture).")
numbered("La cartographie GPEC, qui lit les compétences employé, reflète immédiatement la "
         "progression. La boucle est bouclée.")
h2("10.3 Sourcing & mobilité interne")
para("Pour une offre, le recruteur voit à la fois des candidats externes (matching sur le vivier "
     "global) et des employés internes mobilisables (matching interne + indicateur « formable »), "
     "avec assignation de formations pour combler les écarts. La même offre nourrit donc "
     "recrutement externe et développement interne.")
h2("10.4 GPEC prévisionnelle")
para("Le RH saisit des cibles d'effectif (emploi-type ROME + nombre + horizon) — seule saisie "
     "humaine. Le moteur en déduit, pour chaque cible : qualifiés, potentiels (formables), écart, "
     "compétences émergentes, statut, et une synthèse rédigée par IA (avec repli si l'IA n'est "
     "pas configurée).")

# ══════════════════════════════════════════════════════════════════════════════
# 11. CONCEPTION
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Conception : choix d'ingénierie et justifications")
bullet("Trois apps séparées plutôt qu'un monolithe : le LMS préexistait (template) et a été "
       "intégré par SSO plutôt que réécrit — pragmatique mais source d'hétérogénéité.", "Séparation — ")
bullet("Un seul login (RH) et un secret JWT partagé : SSO simple à mettre en œuvre, expérience "
       "fluide, au prix d'un couplage fort par le secret.", "SSO par secret partagé — ")
bullet("Référentiel ROME comme socle sémantique : permet de raisonner « proximité métier » et "
       "donc potentiel — fondement de la thèse du « savoir-évoluer ».", "ROME — ")
bullet("Matching déterministe et explicable (pondérations fixes, pénalités lisibles) plutôt "
       "qu'un modèle ML opaque : défendable en soutenance (transparence, reproductibilité).", "Déterminisme — ")
bullet("Réutilisation du moteur de matching pour la GPEC (employé=candidat, emploi-type=offre) : "
       "évite la duplication et garantit la cohérence des scores.", "DRY — ")
bullet("Double-écriture candidat+employé à la complétion de cours : assure que la boucle vivante "
       "se reflète dans les deux vues (vivier et organisation).", "Cohérence des profils — ")
bullet("IA en complément (Groq pour générer des offres et la synthèse GPEC ; France Travail pour "
       "enrichir le ROME) avec repli statique : la démo fonctionne sans clés.", "IA optionnelle — ")

# ══════════════════════════════════════════════════════════════════════════════
# 12. CRITIQUE
# ══════════════════════════════════════════════════════════════════════════════
h1("12. CRITIQUE COMPLÈTE — les failles")
para("Section volontairement sans complaisance : ce sont les points à connaître et à assumer "
     "devant le jury. Classés par gravité et par nature.", italic=True, color=RED)

h2("12.1 Sécurité (priorité haute)")
bullet("Le même secret HS256 est codé en dur et committé dans les trois applications "
       "(valeur par défaut « dev_secret_key… »). En production, c'est une faille critique : "
       "quiconque connaît le secret peut forger n'importe quel token, y compris un token "
       "ADMIN.", "Secret JWT partagé et en dur — ")
bullet("Le backend forge lui-même un token de service de rôle ADMIN (sub=\"0\") pour parler au "
       "LMS. Couplé au secret committé, cela signifie qu'un token ADMIN est trivialement "
       "forgeable hors du système.", "Token de service ADMIN — ")
bullet("Le SSO transmet le JWT dans l'URL (?token=). Même s'il est retiré de l'URL côté client "
       "après capture, il a déjà pu transiter dans l'historique, les logs serveur/proxy et "
       "l'en-tête Referer.", "Token dans l'URL — ")
bullet("Le token est stocké dans localStorage (RH et LMS) : vulnérable au vol par XSS, là où un "
       "cookie httpOnly serait plus sûr.", "Stockage localStorage — ")
bullet("Pas de refresh token ni de révocation : un seul access token ; sa fuite est exploitable "
       "jusqu'à expiration.", "Cycle de vie des tokens — ")
bullet("CORS très permissif (regex autorisant tout port localhost + allow_credentials). "
       "Acceptable en dev, à durcir absolument en production.", "CORS — ")
bullet("/api/lms/course-completed fait confiance au LMS sur le niveau de compétence reçu, sans "
       "re-vérifier côté RH la réalité de la réussite : confiance transitive entre services.",
       "Confiance inter-services — ")

h2("12.2 Cohérence multi-tenant")
bullet("Le nom d'entreprise est codé en dur (« TechCorp ») dans le dashboard recruteur et à la "
       "création d'offre, au lieu d'être dérivé du tenant connecté : le multi-tenant est réel en "
       "base mais factice à l'affichage/à la création.", "Entreprise figée — ")
bullet("company_id est injecté dans le JWT via getattr(user, \"company_id\", None) alors que "
       "l'entité User ne porte pas ce champ (c'est Recruiter/Employee qui l'ont) : company_id "
       "vaut donc toujours None dans le token, et côté LMS il est inexploitable.", "company_id absent du token — ")

h2("12.3 Modèle de données & nommage")
bullet("Confusion récurrente employee_id vs user_id : dans le pont LMS et la GPEC, le champ "
       "nommé employeeId/employee_id désigne en réalité l'id utilisateur RH. C'est documenté mais "
       "fragile et propice aux bugs.", "Nommage ambigu — ")
bullet("Historique Alembic avec têtes de fusion (merge heads) et migrations non suivies dans le "
       "répertoire de travail : l'état du schéma n'est pas garanti reproductible.", "Migrations — ")
bullet("MatchResult.has_applied est reconnu redondant avec le statut PENDING et destiné à être "
       "supprimé (dette résiduelle).", "Champ redondant — ")
bullet("Le modèle Evaluation (entretiens annuels) existe mais n'est pas réellement exploité en "
       "interface : fonctionnalité amorcée non terminée.", "Évaluations inertes — ")

h2("12.4 Cohérence produit / LMS")
bullet("L'espace formateur du LMS est en anglais et porte encore le branding du template "
       "d'origine (« Dar Al-Ilm »), alors que tout le reste de la plateforme est en français et "
       "rebrandé « RecruitPRO Academy ». Incohérence très visible en démo.", "Hétérogénéité du LMS — ")
bullet("La sidebar formateur contient des entrées « Students » et « Analytics » qui ne pointent "
       "vers aucune page existante : clics menant à des 404.", "Liens morts — ")
bullet("La page /dashboard du LMS (héritée) lit des champs (firstName, lastName, role "
       "student/instructor/admin) qui ne correspondent pas au mapping réel du JWT RH : affichage "
       "potentiellement cassé.", "Vestige de template — ")
bullet("Deux mécanismes de liaison cours↔compétence coexistent : un lien propre par skillId "
       "(création de cours, GPEC) et une heuristique par correspondance de titre dans la mobilité "
       "du Kanban — le second est fragile.", "Liaison cours/compétence — ")

h2("12.5 Données factices & doublons d'écrans")
bullet("Plusieurs indicateurs sont codés en dur : « Prochaines interviews : 4 », « Visibilité : "
       "Expert », et les chiffres marketing de la landing (« 99.8 % », « +500 entreprises », "
       "« Posté il y a 2j »).", "KPIs en dur — ")
bullet("Le bloc « Mes Formations » du dashboard employé affiche toujours « Aucune formation "
       "assignée » (statique), contredisant la page dédiée qui charge les vraies inscriptions.",
       "Dashboard trompeur — ")
bullet("Deux pages de création d'offre coexistent (/jobs/create avec IA et /jobs/new "
       "formulaire simple) : redondance à clarifier.", "Doublon d'écrans — ")

h2("12.6 Performance & robustesse")
bullet("Le dashboard recruteur charge les matches offre par offre (N+1 requêtes côté client).",
       "N+1 — ")
bullet("Peu de pagination effective : les listes (candidats, employés) sont chargées en bloc.",
       "Pagination — ")
bullet("Gestion d'erreurs inégale : usage d'alert() natifs par endroits, messages génériques "
       "(« Session expirée ? ») ailleurs.", "UX d'erreur — ")

h2("12.7 Pertinence algorithmique (à assumer méthodologiquement)")
bullet("Les pondérations 60/25/15 et la pénalité 0,7^n sont choisies a priori, sans calibration "
       "sur des données réelles ni évaluation de la qualité du matching (pas de vérité terrain).",
       "Paramètres non calibrés — ")
bullet("Le potentiel est opérationnalisé par l'adjacence ROME + un bonus certifications binaire "
       "(+12 si le champ texte est rempli, quel que soit son contenu) : proxy grossier de la "
       "« trainabilité » réelle.", "Potentiel approximé — ")
bullet("Le score de complétude d'onboarding compte deux fois les compétences (≥5 puis >0) et "
       "« bloque » sous 40 % côté interface sans garantie d'application côté backend.", "Score de complétude — ")

h2("12.8 Conformité & exploitation")
bullet("RGPD : CV et données candidats sont conservés et globaux, sans gestion visible du "
       "consentement, de la durée de conservation ou du droit à l'effacement.", "RGPD — ")
bullet("Tests : présence de tests backend (pytest, base recruitment_test) mais couverture réelle "
       "non démontrée ; pas de tests visibles côté frontend / LMS.", "Tests — ")
bullet("Secrets de dev committés (URL DB avec mot de passe, secret JWT) : acceptable pour la "
       "démo, à isoler pour toute mise en production.", "Secrets en dépôt — ")

# ══════════════════════════════════════════════════════════════════════════════
# 13. POINTS FORTS
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Points forts et éléments différenciants")
bullet("Trois domaines RH réunis autour d'un référentiel de compétences unique, avec une vraie "
       "intégration (pas une simple juxtaposition).", "Intégration — ")
bullet("La boucle vivante formation → compétence → GPEC est réellement implémentée (double-"
       "écriture, relève jamais à la baisse) : c'est le cœur différenciant et il fonctionne.", "Boucle vivante — ")
bullet("Le matching à deux axes (fit / potentiel) opérationnalise directement la problématique du "
       "« savoir-évoluer » et reste explicable.", "Double évaluation — ")
bullet("Réutilisation élégante du moteur de matching pour la GPEC (employé=candidat, emploi-type"
       "=offre) : cohérence et non-duplication.", "Conception DRY — ")
bullet("GPEC à deux temps (état des lieux des postes ouverts + prévisionnel par cibles "
       "d'effectif) avec plan d'action directement actionnable.", "GPEC actionnable — ")
bullet("RBAC à deux niveaux (rôle système + rôle interne à permissions par entreprise).", "Sécurité applicative — ")
bullet("Onboarding candidat riche (9 étapes, extraction CV, score de complétude) et promotion "
       "automatique candidat→employé conservant le profil.", "Continuité du profil — ")
bullet("Aide à la création d'offre par ROME et par IA générative, avec traçabilité de la source "
       "des compétences.", "Assistance à la rédaction — ")
bullet("Interface RH soignée et cohérente (design system, thème clair/sombre, sidebar mutualisée).",
       "UX RH — ")

# ══════════════════════════════════════════════════════════════════════════════
# 14. ZONES DE FLOU
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Zones de flou / à clarifier avant la soutenance")
bullet("L'application réelle du seuil de complétude (< 40 % « bloque » la candidature) est-elle "
       "vérifiée côté backend, ou seulement affichée ?")
bullet("Le périmètre multi-entreprise est-il démontrable (créer une 2e entreprise et montrer "
       "l'isolation), ou la démo reste-t-elle mono-tenant « TechCorp » ?")
bullet("Les statistiques formateur du LMS sont-elles correctement scopées par formateur, et "
       "quelles données alimentent réellement totalStudents / totalEnrollments ?")
bullet("Quel est le comportement exact sans clés France Travail / Groq (repli statique) — à "
       "tester avant la démo pour ne pas être surpris.")
bullet("Couverture de tests effective (lancer run_tests.py et noter le pourcentage) pour pouvoir "
       "le citer.")
bullet("Statuts de candidature : vérifier la cohérence de l'enum backend avec les colonnes du "
       "Kanban (notamment SHORTLISTED / « Finalistes »).")

# ══════════════════════════════════════════════════════════════════════════════
# 15. QUESTIONS DU JURY
# ══════════════════════════════════════════════════════════════════════════════
h1("15. Questions probables du jury et éléments de réponse")
def qa(q, a):
    para("Q : " + q, bold=True)
    para("R : " + a)
    spacer()
qa("En quoi est-ce plus qu'un ATS ?",
   "Le système relie recrutement, mobilité et formation autour d'un référentiel unique, et "
   "surtout fait évoluer ce référentiel : la formation réussie relève la compétence et se "
   "répercute dans le pilotage GPEC. C'est un dispositif de gestion de la compétence dans la "
   "durée, pas une simple gestion de candidatures.")
qa("Comment justifiez-vous les pondérations du matching (60/25/15) ?",
   "Ce sont des hypothèses de conception, assumées et explicables, qui privilégient les "
   "compétences. L'algorithme est volontairement déterministe et transparent. La limite — "
   "l'absence de calibration sur données réelles — est identifiée et constitue une piste "
   "d'amélioration (apprentissage des poids sur des recrutements historiques).")
qa("Qu'est-ce que le « potentiel » et pourquoi l'adjacence ROME ?",
   "Le potentiel mesure la capacité à combler les écarts par la formation. Deux métiers proches "
   "au sens ROME partagent des compétences : un écart dans un domaine adjacent est plus "
   "« pontable » qu'un écart dans un domaine étranger. C'est un proxy de la trainabilité, "
   "imparfait mais opérationnel et interprétable.")
qa("Comment fonctionne l'authentification entre les trois applications ?",
   "Un seul login (RH) émet un JWT signé avec un secret partagé ; le LMS reçoit ce token en SSO "
   "(paramètre d'URL puis stockage) et le vérifie avec le même secret. Je connais la limite : en "
   "production il faudrait un secret en variable d'environnement (jamais committé), des cookies "
   "httpOnly et idéalement une signature asymétrique.")
qa("La compétence peut-elle régresser ?",
   "Non : la relève est monotone (jamais à la baisse) à la complétion d'un cours. C'est un choix "
   "assumé pour refléter un acquis ; une gestion de la péremption (last_used existe sur "
   "EmployeeSkill) serait une évolution.")
qa("Le multi-tenant est-il réel ?",
   "Oui au niveau base (company_id sur les tables concernées, filtrage par entreprise du "
   "recruteur), avec candidats et compétences volontairement globaux. À l'affichage, le nom "
   "d'entreprise est encore figé (« TechCorp ») — c'est un point cosmétique à corriger.")

# ══════════════════════════════════════════════════════════════════════════════
# 16. ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
h1("16. Feuille de route d'amélioration")
h2("16.1 Sécurité (avant toute mise en production)")
bullet("Sortir le secret JWT du code (variable d'environnement), le faire tourner, envisager une "
       "signature asymétrique (RS256) et un échange de token côté serveur plutôt que dans l'URL.")
bullet("Cookies httpOnly + SameSite, refresh tokens et révocation ; durcir le CORS.")
bullet("Re-vérifier côté RH la réalité de la complétion de cours (ne pas faire une confiance "
       "aveugle au LMS).")
h2("16.2 Cohérence & finition")
bullet("Dériver le nom d'entreprise du tenant ; corriger company_id dans le token.")
bullet("Homogénéiser le LMS (français, branding « RecruitPRO Academy », supprimer les liens "
       "morts et la page /dashboard héritée).")
bullet("Remplacer les KPIs en dur par de vraies données ; fusionner les deux pages de création "
       "d'offre ; brancher la section formations du dashboard employé.")
h2("16.3 Profondeur fonctionnelle & méthode")
bullet("Calibrer/évaluer le matching (vérité terrain, A/B sur les poids) ; affiner le potentiel "
       "(au-delà du bonus certifications binaire).")
bullet("Activer les évaluations annuelles ; gérer la péremption des compétences (last_used).")
bullet("Conformité RGPD (consentement, conservation, effacement) ; étoffer les tests "
       "(frontend/LMS) et publier la couverture.")

spacer()
para("— Fin du bilan —", italic=True, color=GREY)

out = "BILAN_TECHNIQUE_RecruitPRO_" + datetime.date.today().strftime("%Y-%m-%d") + ".docx"
doc.save(out)
print("OK ->", out)
